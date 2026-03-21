"""
Document creation tools — Word (.docx) and Excel (.xlsx).
Brain calls these via tool system when user asks for reports/documents.
"""
import logging
import os
from pathlib import Path
from datetime import datetime

log = logging.getLogger("seriai.tools.documents")


def create_word_document(title: str, content: str, file_path: str = "") -> dict:
    """
    Create a Word (.docx) document.

    Content format:
        # Heading 1
        ## Heading 2
        ### Heading 3
        - Bullet item
        Normal paragraph text

    Returns dict with 'result' or 'error' key.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "python-docx kütüphanesi yüklü değil. pip install python-docx"}

    try:
        doc = Document()

        # Style defaults
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Title
        doc.add_heading(title, level=0)

        # Metadata line
        doc.add_paragraph(
            f"Oluşturma tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        ).runs[0].font.size = Pt(9)

        # Parse content
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("• "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        # Determine save path
        if not file_path:
            safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)
            file_path = str(Path.home() / "Desktop" / f"{safe_title}.docx")

        # Ensure directory exists
        Path(file_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(file_path)
        log.info(f"Word belgesi oluşturuldu: {file_path}")
        return {"result": f"Belge oluşturuldu: {file_path}", "file_path": file_path}

    except Exception as e:
        log.error(f"Word belgesi oluşturulamadı: {e}")
        return {"error": f"Belge oluşturulamadı: {e}"}


def create_excel_document(title: str, headers: list, rows: list, file_path: str = "") -> dict:
    """
    Create an Excel (.xlsx) document.

    Args:
        title: Sheet title and default filename
        headers: List of column header strings
        rows: List of lists (each inner list is a row of values)
        file_path: Optional save path

    Returns dict with 'result' or 'error' key.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return {"error": "openpyxl kütüphanesi yüklü değil. pip install openpyxl"}

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel sheet name max 31 chars

        # Header style
        header_font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # Write headers
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Write data rows
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="center")

        # Auto-width columns
        for col_idx in range(1, len(headers) + 1):
            col_letter = get_column_letter(col_idx)
            max_width = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 10
            for row_idx in range(2, len(rows) + 2):
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value:
                    max_width = max(max_width, len(str(cell_value)))
            ws.column_dimensions[col_letter].width = min(max_width + 4, 50)

        # Freeze header row
        ws.freeze_panes = "A2"

        # Determine save path
        if not file_path:
            safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)
            file_path = str(Path.home() / "Desktop" / f"{safe_title}.xlsx")

        Path(file_path).parent.mkdir(parents=True, exist_ok=True)
        wb.save(file_path)
        log.info(f"Excel belgesi oluşturuldu: {file_path}")
        return {"result": f"Excel oluşturuldu: {file_path} ({len(rows)} satır)", "file_path": file_path}

    except Exception as e:
        log.error(f"Excel belgesi oluşturulamadı: {e}")
        return {"error": f"Excel oluşturulamadı: {e}"}


def register_document_tools(registry):
    """Register document tools with the tool registry."""
    from seriai.tools.registry import ToolDef

    registry.register(ToolDef(
        name="create_word_document",
        description="Word (.docx) belgesi oluşturur. Markdown benzeri içerik destekler: # başlık, ## alt başlık, - madde.",
        parameters={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Belge başlığı"},
                "content": {"type": "string", "description": "Belge içeriği (Markdown formatında)"},
                "file_path": {"type": "string", "description": "Kayıt yolu (boş bırakılırsa Desktop'a kaydeder)", "default": ""},
            },
            "required": ["title", "content"],
        },
        handler=create_word_document,
        domain="general",
    ))

    registry.register(ToolDef(
        name="create_excel_document",
        description="Excel (.xlsx) belgesi oluşturur. Başlıklar ve satırlar ile tablo formatında.",
        parameters={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Sayfa ve dosya adı"},
                "headers": {"type": "array", "description": "Sütun başlıkları listesi"},
                "rows": {"type": "array", "description": "Veri satırları (liste içinde liste)"},
                "file_path": {"type": "string", "description": "Kayıt yolu (boş bırakılırsa Desktop'a kaydeder)", "default": ""},
            },
            "required": ["title", "headers", "rows"],
        },
        handler=create_excel_document,
        domain="general",
    ))

    log.info("Document tools registered (Word + Excel).")
